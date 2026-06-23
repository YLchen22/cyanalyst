# -*- coding: utf-8 -*-
"""CyAnalyst — 研报 Word 文档生成脚本"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import datetime


# ==================== 颜色常量 ====================
RED_UP = RGBColor(0xDC, 0x26, 0x26)      # 涨 — 红
GREEN_DOWN = RGBColor(0x1A, 0x7A, 0x3A)  # 跌 — 绿
GRAY_FLAT = RGBColor(0x6B, 0x72, 0x80)   # 平
DARK = RGBColor(0x1F, 0x29, 0x37)
ACCENT = RGBColor(0xDC, 0x26, 0x26)


def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def set_narrow_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def add_header_bar(doc, title, subtitle=""):
    """添加文档标题栏（红色主题）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RED_UP
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(subtitle)
        run2.font.size = Pt(12)
        run2.font.color.rgb = GRAY_FLAT
        run2.font.name = 'Microsoft YaHei'
        run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # 分隔线
    doc.add_paragraph('─' * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def add_section_heading(doc, text):
    """添加章节标题"""
    p = doc.add_paragraph()
    run = p.add_run('■ ' + text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RED_UP
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    p.space_after = Pt(6)


def add_paragraph(doc, text, bold=False, color=None, size=10.5):
    """添加正文段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if color:
        run.font.color.rgb = color
    p.paragraph_format.line_spacing = 1.5
    return p


def add_table_with_style(doc, headers, rows, col_widths=None):
    """添加带样式的数据表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        set_cell_shading(cell, 'DC2626')

    # 数据行
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            if r_idx % 2 == 0:
                set_cell_shading(cell, 'F9FAFB')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def add_kpi_box(doc, items):
    """添加关键指标展示"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parts = []
    for label, value, direction in items:
        if direction == 'up':
            color = RED_UP
            arrow = '▲ '
        elif direction == 'down':
            color = GREEN_DOWN
            arrow = '▼ '
        else:
            color = GRAY_FLAT
            arrow = ''

        run = p.add_run(f'{arrow}{label}: {value}    ')
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = color
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.add_paragraph()


def add_disclaimer(doc):
    """添加免责声明"""
    doc.add_paragraph('─' * 60)
    p = doc.add_paragraph()
    run = p.add_run('⚠️ 免责声明：本报告仅供研究参考，不构成任何投资建议。投资有风险，入市需谨慎。')
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY_FLAT
    run.font.italic = True
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    p2 = doc.add_paragraph()
    run2 = p2.add_run(f'Generated by CyAnalyst · {datetime.datetime.now().strftime("%Y-%m-%d %H:%M")}')
    run2.font.size = Pt(8)
    run2.font.color.rgb = GRAY_FLAT
    run2.font.name = 'Microsoft YaHei'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')


# ==================== 简报1：盘前简报 ====================

def generate_morning_brief():
    doc = Document()
    set_narrow_margins(doc)

    add_header_bar(doc, 'CyAnalyst 每日盘前简报', '2026年6月23日（周二）· 盘前预览')

    # --- 一、市场概览 ---
    add_section_heading(doc, '一、前日市场回顾')

    add_paragraph(doc, '昨日（6月22日）A 股三大指数集体收涨，两市成交额 1.08 万亿元，较前一交易日放量约 8%。北向资金全天净流入 52.6 亿元，连续第 5 个交易日净买入。')

    add_table_with_style(doc,
        ['指数', '收盘价', '涨跌幅', '成交额(亿)', '趋势'],
        [
            ['上证指数', '3,287.62', '+1.24%', '4,521', '▲ 5日新高'],
            ['深证成指', '11,023.45', '+1.58%', '5,836', '▲ 站上均线'],
            ['创业板指', '2,291.38', '+2.11%', '2,089', '▲ 量价齐升'],
        ],
        col_widths=[3, 3, 2, 2.5, 3]
    )

    # --- 二、板块表现 ---
    add_section_heading(doc, '二、板块轮动分析')

    add_paragraph(doc, '昨日涨幅居前的板块主要集中在科技成长方向，传统消费和周期板块相对弱势。')
    add_paragraph(doc, '')

    add_table_with_style(doc,
        ['排名', '板块', '涨跌幅', '主力净流入(亿)', '驱动因素'],
        [
            ['1', '半导体', '+4.52%', '+38.2', '国产替代加速 + Q2业绩预期'],
            ['2', 'AI算力', '+3.87%', '+25.6', '算力需求持续超预期'],
            ['3', '新能源车', '+2.94%', '+18.3', '6月销量数据亮眼'],
            ['4', '消费电子', '+2.31%', '+12.1', '新品周期 + 需求回暖'],
            ['5', '食品饮料', '-0.82%', '-8.5', '消费复苏力度不及预期'],
        ],
        col_widths=[1.5, 2.5, 2, 3, 4]
    )

    # --- 三、资金面 ---
    add_section_heading(doc, '三、资金动向')

    add_paragraph(doc, '主力资金持续流入科技板块，北向资金偏好新能源及消费蓝筹。融资余额小幅回升至 1.52 万亿元，市场风险偏好有所修复。')

    add_kpi_box(doc, [
        ('北向净流入', '+52.6亿', 'up'),
        ('主力净流入', '+86.3亿', 'up'),
        ('融资余额', '1.52万亿', 'up'),
        ('两市成交', '1.08万亿', 'up'),
    ])

    # --- 四、今日关注 ---
    add_section_heading(doc, '四、今日交易策略建议')

    add_paragraph(doc, '【短线策略】', bold=True, color=RED_UP)
    add_paragraph(doc, '1. 半导体板块：连续3日放量上涨，短期存在回调压力，建议逢高减仓，等待回调至5日线附近再考虑加仓。')
    add_paragraph(doc, '2. AI算力：趋势向上明确，可沿10日均线逢低布局，关注光模块、服务器等细分方向。')
    add_paragraph(doc, '3. 新能源车：6月销量数据超预期，但板块已累计较大涨幅，建议持有为主，不宜追高。')

    add_paragraph(doc, '【中线布局】', bold=True, color=RED_UP)
    add_paragraph(doc, '1. 消费电子：Q3传统旺季 + 苹果新品周期，可逢低布局产业链核心标的。')
    add_paragraph(doc, '2. 食品饮料：估值已回落至合理区间，关注Q2业绩拐点信号，分批建仓。')

    # --- 五、今日重要事件 ---
    add_section_heading(doc, '五、今日重要事件日历')

    add_table_with_style(doc,
        ['时间', '事件', '影响范围'],
        [
            ['09:30', 'A股开盘', '全市场'],
            ['10:00', '央行公开市场操作', '流动性'],
            ['14:00', '工信部新闻发布会', '科技/制造'],
            ['16:00', '北向资金统计发布', '资金面'],
            ['美股', '隔夜美股：纳指期货+0.3%', 'A股开盘情绪'],
        ],
        col_widths=[2.5, 5, 5]
    )

    add_disclaimer(doc)

    # 保存
    out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'site', 'docs', 'public', 'research', 'daily', '2026-06-23-morning.docx')
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    print(f'[OK] Saved: {out_path}')


# ==================== 简报2：盘后复盘 ====================

def generate_evening_brief():
    doc = Document()
    set_narrow_margins(doc)

    add_header_bar(doc, 'CyAnalyst 每日盘后复盘', '2026年6月22日（周一）· 收盘分析')

    # --- 一、收盘概况 ---
    add_section_heading(doc, '一、今日收盘概况')

    add_paragraph(doc, '今日市场整体呈现"沪弱深强"格局。上证指数全天窄幅震荡，最终微跌；深成指和创业板指在新能源和半导体板块带动下收涨。两市成交额较上一交易日略有萎缩，市场观望情绪渐浓。')

    add_table_with_style(doc,
        ['指数', '收盘', '涨跌', '成交额(亿)', '技术形态'],
        [
            ['上证指数', '3,287.62', '-0.18%', '4,218', '回踩10日线'],
            ['深证成指', '11,023.45', '+0.67%', '5,423', '站上5日线'],
            ['创业板指', '2,291.38', '+1.54%', '1,876', '放量突破'],
            ['科创50', '1,056.83', '+2.12%', '687', '强势突破'],
        ],
        col_widths=[3, 3, 2, 2.5, 3]
    )

    # --- 二、板块深度分析 ---
    add_section_heading(doc, '二、板块涨跌全景')

    add_paragraph(doc, '今日领涨板块集中在半导体、光伏、储能等赛道方向。半导体板块多只个股创出年内新高，市场对Q2业绩预期持续升温。传统消费板块承压，食品饮料板块连续第3日净流出。')

    add_table_with_style(doc,
        ['板块', '涨跌幅', '净流入(亿)', '连涨/跌天数', '驱动逻辑'],
        [
            ['半导体', '+3.52% ▲', '+42.1', '连涨4天', '国产替代 + CIS涨价'],
            ['光伏设备', '+2.87% ▲', '+18.9', '连涨2天', '欧洲需求反弹'],
            ['储能', '+2.45% ▲', '+12.6', '首日反弹', '政策催化'],
            ['白酒', '-1.23% ▼', '-15.7', '连跌3天', '动销数据不及预期'],
            ['银行', '-0.76% ▼', '-9.8', '首日下跌', '降息预期压制'],
        ],
        col_widths=[2.5, 2.5, 2.5, 2.5, 4]
    )

    # --- 三、资金面分析 ---
    add_section_heading(doc, '三、资金面深度分析')

    add_paragraph(doc, '北向资金：今日净流入 52.6 亿元，其中沪股通 18.2 亿元，深股通 34.4 亿元。深股通流入明显大于沪股通，与今日"沪弱深强"格局一致。', size=10.5)
    add_paragraph(doc, '主力资金：全市场净流入 86.3 亿元，其中超大单净流入 62.1 亿元，机构资金参与度较高。', size=10.5)
    add_paragraph(doc, '融资融券：融资余额 1.52 万亿元，较前日增加 28.6 亿元；融券余额持续处于低位，做空力量薄弱。', size=10.5)

    add_kpi_box(doc, [
        ('北向净流入', '+52.6亿', 'up'),
        ('主力净流入', '+86.3亿', 'up'),
        ('融资变动', '+28.6亿', 'up'),
        ('涨跌比', '2,876/1,534', 'up'),
    ])

    # --- 四、重点标的跟踪 ---
    add_section_heading(doc, '四、关注标的池表现')

    add_table_with_style(doc,
        ['代码', '名称', '收盘价', '涨跌', '评分', '操作建议'],
        [
            ['002371', '北方华创', '358.20', '+5.23%', '92 ↑', '持有，逢高可减仓'],
            ['300750', '宁德时代', '228.50', '+2.18%', '85 ↑', '持有'],
            ['688981', '中芯国际', '62.35', '+3.87%', '88 ↑', '加仓关注'],
            ['600519', '贵州茅台', '1,685.00', '-1.02%', '65 ↓', '观望'],
            ['000858', '五粮液', '162.80', '-1.57%', '62 ↓', '减仓'],
        ],
        col_widths=[2, 2.5, 2, 2, 2, 3.5]
    )

    # --- 五、技术面 ---
    add_section_heading(doc, '五、技术面研判')

    add_paragraph(doc, '上证指数日线级别：指数在3,250-3,330区间窄幅震荡已持续12个交易日，MACD红柱缩短，短期面临方向选择。若放量突破3,330点，则打开上行空间；若跌破3,250点支撑，则需警惕回调风险。', size=10.5)

    add_paragraph(doc, '创业板指日线级别：今日放量突破2,250点压力位，MACD金叉初现，短期趋势转强。下一目标位2,350点（前期高点）。', size=10.5)

    add_paragraph(doc, '成交量分析：两市成交额近期在9,500亿至1.1万亿之间波动，未能有效放大，存量博弈特征明显。', size=10.5)

    # --- 六、明日策略 ---
    add_section_heading(doc, '六、明日策略展望')

    add_paragraph(doc, '【核心观点】', bold=True, color=RED_UP)
    add_paragraph(doc, '市场处于方向选择窗口期，维持"震荡偏多"判断。操作上以结构性机会为主，聚焦科技成长板块。')

    add_paragraph(doc, '【策略要点】', bold=True, color=RED_UP)
    add_paragraph(doc, '1. 半导体板块已连续大涨4日，短线进入超买区域，建议控制仓位等待回调。')
    add_paragraph(doc, '2. 新能源板块（光伏+储能）刚启动反弹，仍处于布局窗口期。')
    add_paragraph(doc, '3. 消费板块持续弱势，耐心等待Q2业绩验证，暂不建议左侧抄底。')
    add_paragraph(doc, '4. 整体仓位建议：维持6-7成，留足现金等待方向确认。')

    add_disclaimer(doc)

    out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'site', 'docs', 'public', 'research', 'daily', '2026-06-22-evening.docx')
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    print(f'[OK] Saved: {out_path}')


# ==================== Main ====================

if __name__ == '__main__':
    generate_morning_brief()
    generate_evening_brief()
    print('All reports generated successfully!')
