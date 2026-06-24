# -*- coding: utf-8 -*-
"""生成 3 篇示例研报的 .docx 文件并上传到 Cloudflare R2"""

import os
import sys
import re
sys.path.insert(0, os.path.dirname(__file__))

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 颜色
RED = RGBColor(0xDC, 0x26, 0x26)
GREEN = RGBColor(0x1A, 0x7A, 0x3A)
BLUE = RGBColor(0x25, 0x64, 0xEB)
DARK = RGBColor(0x1E, 0x29, 0x3B)
GRAY = RGBColor(0x6B, 0x72, 0x80)

def set_font(run, size=10.5, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if color:
        run.font.color.rgb = color

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=18, bold=True, color=DARK)
    doc.add_paragraph()  # spacer

def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=11, color=GRAY)
    doc.add_paragraph()

def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=14, bold=True, color=DARK)

def add_subheading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=12, bold=True, color=BLUE)

def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

def add_labeled_para(doc, label, content):
    p = doc.add_paragraph()
    run_bold = p.add_run(label)
    set_font(run_bold, size=10.5, bold=True, color=DARK)
    run_content = p.add_run(content)
    set_font(run_content, size=10.5)

def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        set_font(run_b, size=10.5, bold=True, color=DARK)
    run = p.add_run(text)
    set_font(run, size=10.5)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    # header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            set_font(run, size=10, bold=True)
    # rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                set_font(run, size=9.5)
    doc.add_paragraph()  # spacer

def add_footer_note(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("免责声明：本报告仅供参考，不构成投资建议。")
    set_font(run, size=9, color=GRAY)

# ==================== 研报生成函数 ====================

def generate_daily():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)

    add_title(doc, "2026-06-24 早盘速览")
    add_subtitle(doc, "CyAnalyst 每日研报 | 2026年6月24日")

    add_heading(doc, "市场概况")
    add_body(doc, "今日 A 股三大指数低开高走，沪指收涨 0.42%，深成指涨 0.78%，创业板指涨 1.05%。两市成交额 8,420 亿元，较前一日放量 12%。北向资金全天净流入 52.6 亿元，连续第 5 个交易日净买入。")

    add_heading(doc, "板块动态")

    add_subheading(doc, "AI 算力（+3.2%）")
    add_body(doc, "受国产 AI 芯片量产消息提振，算力板块全天领涨。龙头个股封板，产业链上下游联动明显。光模块、服务器等细分方向资金流入显著。")

    add_subheading(doc, "新能源（+1.8%）")
    add_body(doc, "经过前期调整，新能源板块企稳反弹。锂电池细分方向资金回流显著，6 月销量数据超预期提振市场信心。")

    add_subheading(doc, "消费电子（+1.5%）")
    add_body(doc, "Q3 传统旺季叠加苹果新品周期预期，消费电子板块温和上行，产业链核心标的获资金关注。")

    add_subheading(doc, "消费（-0.5%）")
    add_body(doc, "消费板块小幅回调，家电领跌。市场风格仍偏成长，资金更青睐科技方向。")

    add_heading(doc, "资金流向")
    add_table(doc, ["方向", "净流入（亿元）"], [
        ["AI 算力", "+68.2"],
        ["新能源", "+32.5"],
        ["消费电子", "+18.6"],
        ["消费", "-12.1"],
        ["金融", "-8.3"],
    ])

    add_heading(doc, "今日策略建议")
    add_subheading(doc, "短线策略")
    add_bullet(doc, "AI 算力：趋势向上明确，可沿 10 日均线逢低布局，关注光模块、服务器等细分方向。", "1. ")
    add_bullet(doc, "新能源：6 月销量数据超预期，但板块已累计较大涨幅，建议持有为主，不宜追高。", "2. ")
    add_bullet(doc, "消费电子：Q3 传统旺季 + 苹果新品周期，可逢低布局产业链核心标的。", "3. ")
    add_subheading(doc, "中线布局")
    add_bullet(doc, "食品饮料：估值已回落至合理区间，关注 Q2 业绩拐点信号，分批建仓。", "1. ")
    add_bullet(doc, "半导体：国产替代加速 + CIS 涨价，中期逻辑不变，回调即机会。", "2. ")

    add_heading(doc, "后市展望")
    add_body(doc, "短期市场风险偏好修复，科技成长主线有望延续。关注 AI 应用落地节奏和半年报业绩验证窗口。")

    add_footer_note(doc)

    path = os.path.join(os.path.dirname(__file__), 'output', 'research', 'daily', '2026-06-24-morning.docx')
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)
    return path


def generate_weekly():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)

    add_title(doc, "第25周市场回顾：科技主线回归")
    add_subtitle(doc, "CyAnalyst 周度研报 | 2026年6月15日 — 6月21日")

    add_heading(doc, "本周回顾")
    add_body(doc, "本周（6/17-6/21）A 股缩量反弹，沪指周涨 1.2%，深成指周涨 2.1%，创业板指周涨 2.8%。市场风格明显转向成长，科技板块成为最大赢家。")

    add_heading(doc, "核心数据")
    add_table(doc, ["指数", "周涨跌", "周均成交额（亿）"], [
        ["沪指", "+1.2%", "3,850"],
        ["深成指", "+2.1%", "4,420"],
        ["创业板指", "+2.8%", "1,680"],
    ])

    add_heading(doc, "主线分析")

    add_subheading(doc, "AI 算力产业链")
    add_body(doc, "本周 AI 算力板块涨幅达 5.8%，连续三周跑赢大盘。")
    add_body(doc, "核心驱动因素：")
    add_bullet(doc, "国产 AI 芯片量产进度超预期", "1. ")
    add_bullet(doc, "云厂商资本开支上调，算力需求确定性增强", "2. ")
    add_bullet(doc, "应用层爆发带动推理算力需求阶梯式上升", "3. ")

    add_subheading(doc, "半导体设备")
    add_body(doc, "半导体设备板块周涨 4.2%，与 AI 算力形成共振。国产替代逻辑持续兑现，设备招标节奏加快。")

    add_heading(doc, "下周关注")
    add_bullet(doc, "美联储议息会议（6/25）", "• ")
    add_bullet(doc, "国内 6 月 PMI 数据（6/30）", "• ")
    add_bullet(doc, "AI 应用层半年报业绩预告窗口", "• ")

    add_footer_note(doc)

    path = os.path.join(os.path.dirname(__file__), 'output', 'research', 'weekly', '2026-06-21-week-25.docx')
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)
    return path


def generate_special():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)

    add_title(doc, "AI 芯片国产替代：从0到1的突破")
    add_subtitle(doc, "CyAnalyst 深度专题 | 2026年6月15日")

    add_heading(doc, "核心观点")
    add_body(doc, '国产 AI 芯片产业正在经历从\u201c可用\u201d到\u201c好用\u201d的关键跨越。随着制程瓶颈突破和生态成熟，2026 年有望成为国产 AI 芯片规模化落地的元年。')

    add_heading(doc, "产业背景")
    add_subheading(doc, "需求端：算力缺口持续扩大")
    add_body(doc, '全球 AI 算力需求每 3-4 个月翻一番，而供给端受制于先进制程限制，缺口持续扩大。国产替代从\u201c可选\u201d变为\u201c必选\u201d。')

    add_subheading(doc, "供给端：制程与封装双突破")
    add_bullet(doc, "7nm 量产出货，5nm 进入风险量产", "制程：")
    add_bullet(doc, "2.5D/3D 封装产能爬坡，CoWoS 国产替代加速", "先进封装：")
    add_bullet(doc, "CUDA 兼容层成熟度提升，迁移成本下降", "生态：")

    add_heading(doc, "产业链图谱")
    add_body(doc, "设计 → 制造 → 封装 → 应用")
    add_body(doc, "芯片设计 / 晶圆代工 / 先进封装 / 云计算")
    add_body(doc, "EDA工具 / 光刻设备 / 测试设备 / AI训练")
    add_body(doc, "IP授权 / 刻蚀设备 / 板级封装 / AI推理")

    add_heading(doc, "投资逻辑")
    add_subheading(doc, "第一层：芯片设计")
    add_body(doc, "直接受益于国产替代浪潮。关注产品力已达商用门槛、客户验证进展顺利的设计公司。")
    add_subheading(doc, "第二层：设备材料")
    add_body(doc, "国产化率最低、弹性最大的环节。光刻、刻蚀、薄膜沉积等关键设备国产化率均低于 20%。")
    add_subheading(doc, "第三层：先进封装")
    add_body(doc, "AI 芯片对先进封装依赖度极高，2.5D/3D 封装产能成为产业链瓶颈。国内封测龙头已具备量产能力。")

    add_heading(doc, "风险提示")
    add_bullet(doc, "技术迭代不及预期")
    add_bullet(doc, "下游需求波动")
    add_bullet(doc, "地缘政策变化")
    add_bullet(doc, "估值溢价过高")

    add_heading(doc, "结论")
    add_body(doc, 'AI 芯片国产替代是未来 3 年确定性最高的产业趋势之一。建议沿着\u201c设计-设备-封装\u201d三层逻辑，寻找具有核心技术壁垒和客户验证进展的标的。')

    add_footer_note(doc)

    path = os.path.join(os.path.dirname(__file__), 'output', 'research', 'special', '2026-06-15-ai-chips.docx')
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)
    return path


# ==================== 主流程 ====================

def main():
    print("=== Generating Sample Research Reports ===")
    print()
    
    daily_path = generate_daily()
    print(f"[OK] Daily:  {daily_path}")
    
    weekly_path = generate_weekly()
    print(f"[OK] Weekly: {weekly_path}")
    
    special_path = generate_special()
    print(f"[OK] Special: {special_path}")
    
    print()
    print("All done. Now run upload_to_r2.py --dir scripts/output '' to upload.")
    
    return [daily_path, weekly_path, special_path]


if __name__ == "__main__":
    main()
